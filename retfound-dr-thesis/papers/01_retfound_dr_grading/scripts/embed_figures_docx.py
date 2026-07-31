"""Generate thesis bar charts with Pillow (no matplotlib) and embed in Word."""
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "figures"
FIG.mkdir(exist_ok=True)
DOCX = ROOT / "manuscript" / "Thesis_Report_RETFound_Baseline_vs_Enhanced.docx"

# Final-run numbers
APTOS = {
    "Baseline": {"Accuracy": 0.831, "Macro-F1": 0.671, "QWK": 0.888, "Ref Acc": 0.929, "Ref AUROC": 0.981},
    "Enhanced": {"Accuracy": 0.811, "Macro-F1": 0.683, "QWK": 0.893, "Ref Acc": 0.935, "Ref AUROC": 0.987},
}
MESS = {
    "Baseline": {"Accuracy": 0.612, "Macro-F1": 0.330, "QWK": 0.488, "Ref Acc": 0.807, "Ref AUROC": 0.725},
    "Enhanced": {"Accuracy": 0.593, "Macro-F1": 0.349, "QWK": 0.505, "Ref Acc": 0.802, "Ref AUROC": 0.778},
}
ABL = {
    "Baseline": {"QWK": 0.888, "Ref AUROC": 0.981, "Accuracy": 0.831},
    "Enhanced": {"QWK": 0.893, "Ref AUROC": 0.987, "Accuracy": 0.811},
    "A1": {"QWK": 0.892, "Ref AUROC": 0.985, "Accuracy": 0.738},
    "A2": {"QWK": 0.895, "Ref AUROC": 0.987, "Accuracy": 0.765},
    "A3": {"QWK": 0.881, "Ref AUROC": 0.988, "Accuracy": 0.764},
}

COLORS = {
    "Baseline": (31, 78, 121),
    "Enhanced": (196, 89, 17),
    "A1": (84, 130, 53),
    "A2": (112, 48, 160),
    "A3": (131, 60, 12),
}


def font(size=16, bold=False):
    candidates = [
        r"C:\Windows\Fonts\timesbd.ttf" if bold else r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
    ]
    for p in candidates:
        if Path(p).exists():
            return ImageFont.truetype(p, size)
    return ImageFont.load_default()


def draw_grouped_bars(path, title, series, metrics, ymax=1.05):
    """series: {model_name: {metric: value}}"""
    W, H = 1100, 620
    img = Image.new("RGB", (W, H), (255, 255, 255))
    d = ImageDraw.Draw(img)
    left, right, top, bottom = 90, 40, 70, 110
    plot_w, plot_h = W - left - right, H - top - bottom

    d.text((W // 2, 18), title, fill=(20, 20, 20), font=font(22, True), anchor="ma")

    # axes
    d.line([(left, top), (left, top + plot_h), (left + plot_w, top + plot_h)], fill=(60, 60, 60), width=2)

    # y ticks
    for i in range(6):
        yv = i / 5 * ymax
        y = top + plot_h - (yv / ymax) * plot_h
        d.line([(left - 5, y), (left + plot_w, y)], fill=(230, 230, 230), width=1)
        d.text((left - 10, y), f"{yv:.1f}", fill=(80, 80, 80), font=font(12), anchor="rm")

    models = list(series.keys())
    n_m, n_g = len(models), len(metrics)
    group_w = plot_w / n_g
    bar_w = group_w / (n_m + 1.2)

    for gi, metric in enumerate(metrics):
        gx = left + gi * group_w + group_w * 0.15
        for mi, model in enumerate(models):
            val = series[model][metric]
            bh = (val / ymax) * plot_h
            x0 = gx + mi * bar_w
            y0 = top + plot_h - bh
            color = COLORS.get(model, (100, 100, 100))
            d.rectangle([x0, y0, x0 + bar_w * 0.9, top + plot_h], fill=color)
            d.text((x0 + bar_w * 0.45, y0 - 8), f"{val:.3f}", fill=(30, 30, 30), font=font(11), anchor="mb")
        d.text((left + gi * group_w + group_w / 2, top + plot_h + 18), metric,
               fill=(30, 30, 30), font=font(13), anchor="mt")

    # legend
    lx, ly = left, H - 42
    for model in models:
        color = COLORS.get(model, (100, 100, 100))
        d.rectangle([lx, ly, lx + 18, ly + 18], fill=color)
        d.text((lx + 24, ly + 9), model, fill=(20, 20, 20), font=font(13), anchor="lm")
        lx += 24 + int(d.textlength(model, font=font(13))) + 28

    img.save(path)
    print("saved", path.name)


def draw_dual_dataset(path):
    W, H = 1000, 480
    img = Image.new("RGB", (W, H), (255, 255, 255))
    d = ImageDraw.Draw(img)
    d.text((W // 2, 16), "Internal vs external generalization", fill=(20, 20, 20), font=font(22, True), anchor="ma")

    panels = [
        ("Quadratic Weighted Kappa", "QWK", [(0.888, 0.893), (0.488, 0.505)]),
        ("Referable AUROC", "Ref AUROC", [(0.981, 0.987), (0.725, 0.778)]),
    ]
    sets = ["APTOS", "Messidor-2"]
    for pi, (title, _, vals) in enumerate(panels):
        left = 70 + pi * 480
        top, pw, ph = 70, 400, 300
        d.text((left + pw / 2, 48), title, fill=(20, 20, 20), font=font(16, True), anchor="ma")
        d.line([(left, top), (left, top + ph), (left + pw, top + ph)], fill=(60, 60, 60), width=2)
        for i in range(6):
            yv = i / 5
            y = top + ph - yv * ph
            d.line([(left, y), (left + pw, y)], fill=(235, 235, 235), width=1)
            d.text((left - 8, y), f"{yv:.1f}", fill=(80, 80, 80), font=font(11), anchor="rm")
        group_w = pw / 2
        bar_w = group_w * 0.28
        for gi, sname in enumerate(sets):
            b, e = vals[gi]
            cx = left + gi * group_w + group_w / 2
            for vi, (val, model) in enumerate([(b, "Baseline"), (e, "Enhanced")]):
                bh = val * ph
                x0 = cx - bar_w - 4 if vi == 0 else cx + 4
                y0 = top + ph - bh
                d.rectangle([x0, y0, x0 + bar_w, top + ph], fill=COLORS[model])
                d.text((x0 + bar_w / 2, y0 - 6), f"{val:.3f}", fill=(20, 20, 20), font=font(11), anchor="mb")
            d.text((cx, top + ph + 14), sname, fill=(20, 20, 20), font=font(14), anchor="mt")

    # legend
    d.rectangle([320, H - 40, 338, H - 22], fill=COLORS["Baseline"])
    d.text((346, H - 31), "RETFound Baseline", fill=(20, 20, 20), font=font(13), anchor="lm")
    d.rectangle([540, H - 40, 558, H - 22], fill=COLORS["Enhanced"])
    d.text((566, H - 31), "Enhanced RETFound", fill=(20, 20, 20), font=font(13), anchor="lm")
    img.save(path)
    print("saved", path.name)


def draw_ablation(path):
    W, H = 1100, 520
    img = Image.new("RGB", (W, H), (255, 255, 255))
    d = ImageDraw.Draw(img)
    d.text((W // 2, 16), "Ablation study on APTOS test", fill=(20, 20, 20), font=font(22, True), anchor="ma")

    names = ["Baseline", "Enhanced", "A1", "A2", "A3"]
    labels = ["Baseline", "Enhanced", "A1: LoRA+focal", "A2: +multi-scale", "A3: +ord/ref"]
    panels = [
        ("QWK", "QWK", (0.86, 0.91)),
        ("Referable AUROC", "Ref AUROC", (0.975, 0.995)),
    ]
    for pi, (title, key, (ymin, ymax)) in enumerate(panels):
        left = 60 + pi * 540
        top, pw, ph = 70, 480, 320
        d.text((left + pw / 2, 48), title, fill=(20, 20, 20), font=font(16, True), anchor="ma")
        d.line([(left, top), (left, top + ph), (left + pw, top + ph)], fill=(60, 60, 60), width=2)
        for i in range(6):
            yv = ymin + (ymax - ymin) * i / 5
            y = top + ph - (i / 5) * ph
            d.line([(left, y), (left + pw, y)], fill=(235, 235, 235), width=1)
            d.text((left - 8, y), f"{yv:.3f}", fill=(80, 80, 80), font=font(10), anchor="rm")
        bw = pw / (len(names) + 0.8)
        for i, name in enumerate(names):
            val = ABL[name][key]
            bh = ((val - ymin) / (ymax - ymin)) * ph
            x0 = left + 20 + i * bw
            y0 = top + ph - bh
            d.rectangle([x0, y0, x0 + bw * 0.75, top + ph], fill=COLORS[name])
            d.text((x0 + bw * 0.375, y0 - 6), f"{val:.3f}", fill=(20, 20, 20), font=font(11), anchor="mb")
            d.text((x0 + bw * 0.375, top + ph + 10), labels[i], fill=(20, 20, 20), font=font(11), anchor="mt")
    img.save(path)
    print("saved", path.name)


def draw_acc_qwk(path):
    W, H = 1000, 520
    img = Image.new("RGB", (W, H), (255, 255, 255))
    d = ImageDraw.Draw(img)
    d.text((W // 2, 16), "Ablation trade-off: Accuracy vs QWK (APTOS)", fill=(20, 20, 20),
           font=font(20, True), anchor="ma")
    names = ["Baseline", "Enhanced", "A1", "A2", "A3"]
    labels = ["Baseline", "Enhanced", "A1", "A2", "A3"]
    left, top, pw, ph = 80, 70, 860, 340
    ymin, ymax = 0.70, 0.92
    d.line([(left, top), (left, top + ph), (left + pw, top + ph)], fill=(60, 60, 60), width=2)
    for i in range(6):
        yv = ymin + (ymax - ymin) * i / 5
        y = top + ph - (i / 5) * ph
        d.line([(left, y), (left + pw, y)], fill=(235, 235, 235), width=1)
        d.text((left - 8, y), f"{yv:.2f}", fill=(80, 80, 80), font=font(11), anchor="rm")
    group = pw / len(names)
    bw = group * 0.28
    c_acc, c_qwk = (91, 155, 213), (237, 125, 49)
    for i, name in enumerate(names):
        cx = left + i * group + group / 2
        for val, color, shift in [
            (ABL[name]["Accuracy"], c_acc, -bw - 3),
            (ABL[name]["QWK"], c_qwk, 3),
        ]:
            bh = ((val - ymin) / (ymax - ymin)) * ph
            x0 = cx + shift
            y0 = top + ph - bh
            d.rectangle([x0, y0, x0 + bw, top + ph], fill=color)
            d.text((x0 + bw / 2, y0 - 5), f"{val:.3f}", fill=(20, 20, 20), font=font(10), anchor="mb")
        d.text((cx, top + ph + 14), labels[i], fill=(20, 20, 20), font=font(14), anchor="mt")
    d.rectangle([320, H - 40, 338, H - 22], fill=c_acc)
    d.text((346, H - 31), "Accuracy", fill=(20, 20, 20), font=font(13), anchor="lm")
    d.rectangle([470, H - 40, 488, H - 22], fill=c_qwk)
    d.text((496, H - 31), "QWK", fill=(20, 20, 20), font=font(13), anchor="lm")
    img.save(path)
    print("saved", path.name)


# Generate figures
draw_grouped_bars(
    FIG / "fig1_aptos_comparison.png",
    "APTOS test performance",
    APTOS,
    ["Accuracy", "Macro-F1", "QWK", "Ref Acc", "Ref AUROC"],
)
draw_grouped_bars(
    FIG / "fig2_messidor_comparison.png",
    "Messidor-2 external test performance",
    MESS,
    ["Accuracy", "Macro-F1", "QWK", "Ref Acc", "Ref AUROC"],
)
draw_dual_dataset(FIG / "fig3_internal_vs_external.png")
draw_ablation(FIG / "fig4_ablation_qwk_auroc.png")
draw_acc_qwk(FIG / "fig5_ablation_acc_vs_qwk.png")

# -------- Word doc --------

def set_run_font(run, bold=False, italic=False, size=12):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_mixed(doc, text, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2]); set_run_font(run, bold=True)
        else:
            run = p.add_run(part); set_run_font(run)


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, size=16 if level == 1 else 13)


def shade(cell, color="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=9)
        shade(cell)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=9)
    doc.add_paragraph()


def add_figure(doc, path, caption, width=6.1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, italic=True, size=10)
    cap.paragraph_format.space_after = Pt(14)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run)


doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("Enhancing RETFound for Diabetic Retinopathy Grading:\nBaseline vs Clinically Adapted Fine-Tuning")
set_run_font(run, bold=True, size=16)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = s.add_run("Thesis Experimental Report (APTOS + Messidor-2 + Ablations)")
set_run_font(run, italic=True, size=12)
doc.add_paragraph()

add_h(doc, "Abstract")
add_mixed(
    doc,
    "We compare **RETFound Baseline** with **Enhanced RETFound** for five-class DR grading. "
    "On APTOS, Enhanced improved QWK (0.893 vs 0.888) and referable AUROC (0.987 vs 0.981); "
    "Baseline kept higher accuracy (0.831 vs 0.811). On Messidor-2, Enhanced improved QWK "
    "(0.505 vs 0.488) and referable AUROC (0.778 vs 0.725). Ablation A2 (LoRA + multi-scale + focal) "
    "achieved the best APTOS QWK (0.895)."
)

add_h(doc, "1. Introduction")
add_mixed(
    doc,
    "Diabetic retinopathy grading is ordinal and class-imbalanced. This study asks whether "
    "LoRA + multi-scale + ordinal/focal adaptation of RETFound improves grading over standard fine-tuning, "
    "including under external domain shift."
)

add_h(doc, "2. Methods")
add_table(doc, ["Dataset", "Role"], [
    ["APTOS 2019", "Train/val/test; test N=550"],
    ["Messidor-2", "External test only"],
])
add_table(doc, ["Model", "Description"], [
    ["RETFound Baseline", "Full fine-tuning + weighted CE"],
    ["Enhanced RETFound", "LoRA + multi-scale + focal/ordinal/referable"],
    ["A1 / A2 / A3", "Ablations of Enhanced components"],
])

add_h(doc, "3. Results")
add_h(doc, "3.1 Overall performance", 2)
add_table(doc, ["Model", "Set", "Acc", "Macro-F1", "QWK", "Ref Acc", "Ref AUROC"], [
    ["RETFound Baseline", "APTOS", "0.831", "0.671", "0.888", "0.929", "0.981"],
    ["Enhanced RETFound", "APTOS", "0.811", "0.683", "0.893", "0.935", "0.987"],
    ["RETFound Baseline", "Messidor-2", "0.612", "0.330", "0.488", "0.807", "0.725"],
    ["Enhanced RETFound", "Messidor-2", "0.593", "0.349", "0.505", "0.802", "0.778"],
])

add_figure(doc, FIG / "fig1_aptos_comparison.png",
           "Figure 1. APTOS held-out test metrics: RETFound Baseline vs Enhanced RETFound.")
add_figure(doc, FIG / "fig2_messidor_comparison.png",
           "Figure 2. Messidor-2 external test metrics: RETFound Baseline vs Enhanced RETFound.")
add_figure(doc, FIG / "fig3_internal_vs_external.png",
           "Figure 3. Internal vs external QWK and referable AUROC. Enhanced RETFound shows a clearer "
           "external gain in referable discrimination (0.725 → 0.778).")

bullets(doc, [
    "On APTOS, Enhanced slightly improves ordinal/screening metrics; Baseline wins raw accuracy.",
    "On Messidor-2, Enhanced improves QWK (+0.017) and referable AUROC (+0.053).",
])

add_h(doc, "3.2 Ablation study", 2)
add_table(doc, ["Variant", "Acc", "Macro-F1", "QWK", "Ref AUROC"], [
    ["RETFound Baseline", "0.831", "0.671", "0.888", "0.981"],
    ["Enhanced RETFound", "0.811", "0.683", "0.893", "0.987"],
    ["A1 LoRA + late + focal", "0.738", "0.613", "0.892", "0.985"],
    ["A2 LoRA + multi-scale + focal", "0.765", "0.632", "0.895", "0.987"],
    ["A3 + ordinal + referable", "0.764", "0.629", "0.881", "0.988"],
])
add_figure(doc, FIG / "fig4_ablation_qwk_auroc.png",
           "Figure 4. Ablation QWK (left) and referable AUROC (right) on APTOS. A2 best QWK; A3 best referable AUROC.")
add_figure(doc, FIG / "fig5_ablation_acc_vs_qwk.png",
           "Figure 5. Accuracy vs QWK trade-off across ablation variants.")

bullets(doc, [
    "A1 shows LoRA+focal already nearly matches Enhanced QWK.",
    "A2 (multi-scale) gives the best grading QWK (0.895).",
    "A3 boosts referable AUROC but lowers QWK under the short ablation schedule.",
])

add_h(doc, "3.3 Kaggle figure checklist (paste if available)", 2)
add_mixed(
    doc,
    "Also insert from Kaggle `outputs/figures/` if downloaded: confusion matrices "
    "(`cm_baseline_aptos.png`, `cm_enhanced_aptos.png`, `cm_*_external.png`), "
    "ROC curves (`roc_referable_aptos.png`, `roc_referable_external.png`), "
    "and train/val loss–QWK curves for Baseline/Enhanced/A1–A3."
)

add_h(doc, "4. Discussion")
for i, item in enumerate([
    "Baseline is already strong on APTOS.",
    "Enhanced is comparable internally and stronger externally on QWK/referable AUROC.",
    "A2 is the best pure grading recipe; full Enhanced suits referable screening.",
    "Report both accuracy and ordinal/screening metrics.",
], 1):
    p = doc.add_paragraph(f"{i}. {item}")
    for run in p.runs:
        set_run_font(run)

add_h(doc, "5. Conclusion")
add_mixed(
    doc,
    "Enhanced RETFound improves ordinal/screening behaviour and Messidor-2 referable discrimination "
    "relative to standard RETFound fine-tuning. Multi-scale LoRA fine-tuning is the main useful enhancement."
)

doc.save(DOCX)
print("Updated:", DOCX)
