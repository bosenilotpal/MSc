from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
out_path = ROOT / "manuscript" / "Cover_Letter_RETFound_DR.docx"

TITLE = (
    "LoRA-Adapted RETFound with Multi-Scale Fusion and Ordinal Losses "
    "for Five-Class Diabetic Retinopathy Grading and External Validation"
)

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.paragraph_format.space_after = Pt(10)
style.paragraph_format.line_spacing = 1.15


def set_run_font(run, bold=False, italic=False, size=12):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def para(text="", bold=False, italic=False, size=12, space_after=10,
         align=None, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    if text:
        run = p.add_run(text)
        set_run_font(run, bold=bold, italic=italic, size=size)
    return p


def bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        set_run_font(run)
    p.paragraph_format.space_after = Pt(4)
    return p


# --- Sender / date ---
para("Nilotpal Bose", bold=True, space_after=0)
para("Department of Computer Science and Engineering", space_after=0)
para("Email: [update-email]", space_after=0)
para(date.today().strftime("%B %d, %Y"), space_after=16)

# --- Addressee ---
para("The Editor-in-Chief", space_after=0)
para("[Journal Name]", italic=True, space_after=16)

# --- Salutation ---
para("Dear Editor,", space_after=12)

# --- Opening ---
para(
    "I am pleased to submit our original research manuscript entitled "
    f"\u201c{TITLE}\u201d for consideration for publication in your journal. "
    "The work addresses a practical and clinically important question: how best "
    "to adapt a retinal foundation model (RETFound) for ordered, class-imbalanced "
    "diabetic retinopathy (DR) severity grading, and whether a parameter-efficient, "
    "clinically motivated adaptation generalizes better across datasets than "
    "standard full fine-tuning."
)

# --- Background / motivation ---
para(
    "Automated DR grading is typically framed as a flat five-class classification "
    "problem, which ignores the ordinal structure of the ICDR scale (grades 0\u20134) "
    "and the severe class imbalance of screening data. Foundation models such as "
    "RETFound offer strong retinal representations, but the optimal recipe for "
    "adapting them to DR grading, and the extent to which such gains hold under "
    "external domain shift, remain open questions. Our study directly compares two "
    "adaptation strategies under identical backbones and preprocessing, and evaluates "
    "them both in-domain and on a fully external test set."
)

# --- Contributions ---
para("The principal contributions of this work are:", space_after=6)
bullet(
    "A controlled comparison of RETFound Baseline (full fine-tuning with weighted "
    "cross-entropy) against Enhanced RETFound (LoRA adapters on attention qkv with "
    "~1.3% trainable parameters, multi-scale token fusion from blocks 7/15/23, and "
    "focal + CORAL ordinal + referable auxiliary losses)."
)
bullet(
    "In-domain evaluation on APTOS 2019 (test N = 550), where Enhanced RETFound "
    "slightly improves quadratic weighted kappa (QWK 0.893 vs 0.888) and referable "
    "AUROC (0.987 vs 0.981), while the baseline retains higher overall accuracy "
    "(0.831 vs 0.811)."
)
bullet(
    "External validation on Messidor-2 without any fine-tuning, where Enhanced "
    "RETFound shows clearer, more clinically meaningful gains in QWK (0.505 vs 0.488) "
    "and referable AUROC (0.778 vs 0.725)."
)
bullet(
    "An ablation study isolating each component, showing that LoRA + multi-scale + "
    "focal (A2) achieves the best APTOS QWK (0.895), while ordinal/referable "
    "auxiliaries primarily boost screening-oriented AUROC."
)

# --- Significance ---
para(
    "Taken together, our findings argue that foundation-model adaptation for DR should "
    "be assessed using clinically meaningful, ordinal and screening-oriented endpoints "
    "with external validation, rather than in-domain accuracy alone. We believe these "
    "results are relevant to your readership working on medical image analysis, "
    "foundation-model adaptation, and diabetic retinopathy screening."
)

# --- Standard declarations ---
para(
    "We confirm that this manuscript is original, has not been published previously, "
    "and is not under consideration for publication elsewhere. All authors have "
    "approved the manuscript and agree to its submission. The study uses publicly "
    "available, de-identified datasets (APTOS 2019 and Messidor-2), and the authors "
    "declare no conflicts of interest."
)

para(
    "Thank you for your time and consideration. We look forward to your response and "
    "would be happy to address any questions regarding the manuscript."
)

# --- Sign-off ---
para("Sincerely,", space_after=16)
para("Nilotpal Bose", bold=True, space_after=0)
para("Department of Computer Science and Engineering", space_after=0)
para("Email: [update-email]", space_after=0)

doc.save(out_path)
print("Saved:", out_path)
