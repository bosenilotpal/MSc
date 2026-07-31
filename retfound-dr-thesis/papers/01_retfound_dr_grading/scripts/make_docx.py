from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

ROOT = Path(__file__).resolve().parents[1]
out_path = ROOT / "manuscript" / "Thesis_Report_RETFound_Baseline_vs_Enhanced.docx"
doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.15


def set_run_font(run, bold=False, italic=False, size=12):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_mixed(text, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2]); set_run_font(run, bold=True)
        else:
            run = p.add_run(part); set_run_font(run)
    return p


def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, size=16 if level == 1 else 13)
    return h


def shade(cell, color="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=9)
        shade(cell)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            clean = re.sub(r"\*\*(.*?)\*\*", r"\1", str(val))
            run = cell.paragraphs[0].add_run(clean)
            set_run_font(run, bold=("**" in str(val)), size=9)
    doc.add_paragraph()


def bullets(items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run)


# Title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("Enhancing RETFound for Diabetic Retinopathy Grading:\nBaseline vs Clinically Adapted Fine-Tuning")
set_run_font(run, bold=True, size=16)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = s.add_run("Thesis Experimental Report (APTOS + Messidor-2 + Ablations)")
set_run_font(run, italic=True, size=12)
doc.add_paragraph()

add_h("Abstract")
add_mixed(
    "We compare **RETFound Baseline** (full fine-tuning + cross-entropy) with **Enhanced RETFound** "
    "(LoRA, multi-scale fusion, and focal/ordinal/referable losses) for five-class DR grading. On APTOS 2019 "
    "(test N = 550), Enhanced RETFound slightly improved QWK (0.893 vs 0.888) and referable AUROC "
    "(0.987 vs 0.981), while the baseline kept higher accuracy (0.831 vs 0.811). On external Messidor-2, "
    "Enhanced RETFound improved QWK (0.505 vs 0.488) and referable AUROC (0.778 vs 0.725) despite lower "
    "accuracy. Ablations showed LoRA + multi-scale + focal (A2) achieved the best APTOS QWK (0.895); adding "
    "ordinal/referable losses raised referable AUROC but reduced QWK under a short training schedule. "
    "Enhanced RETFound is comparable internally and stronger on external referable discrimination."
)

add_h("1. Introduction")
add_mixed(
    "Diabetic retinopathy (DR) grading is **ordinal** (ICDR 0–4) and **class-imbalanced**. RETFound provides "
    "strong retinal representations, but standard full fine-tuning with flat cross-entropy may under-emphasise "
    "rare grades and ordinal structure."
)
add_mixed("Research question:")
q = doc.add_paragraph()
q.paragraph_format.left_indent = Inches(0.25)
run = q.add_run(
    "Does a clinically motivated adaptation of RETFound (LoRA + multi-scale fusion + ordinal/focal learning) "
    "improve DR grading relative to standard RETFound fine-tuning, including under external domain shift?"
)
set_run_font(run, italic=True)

add_h("2. Methods")
add_h("2.1 Datasets", 2)
add_table(
    ["Dataset", "Role"],
    [
        ["APTOS 2019", "Train/val/test (70/15/15; fixed splits.json); test N = 550"],
        ["Messidor-2", "External test only (no fine-tuning)"],
    ],
)
bullets([
    "Tasks: five-class ICDR grading; secondary endpoint referable DR (grade ≥ 2).",
    "Input: 224×224 fundus images; ImageNet normalisation; light train-time augmentation.",
])

add_h("2.2 Models", 2)
add_table(
    ["Model", "Description"],
    [
        ["RETFound Baseline", "RETFound ViT-L/16 (CFP weights); full fine-tuning; weighted cross-entropy"],
        ["Enhanced RETFound", "Frozen backbone + LoRA on qkv (~1.3% trainable); multi-scale tokens from blocks 7/15/23; focal + CORAL ordinal + referable auxiliary losses"],
    ],
)
add_mixed("Checkpoints selected by **best validation QWK**.")

add_h("2.3 Ablations (APTOS, 15 epochs)", 2)
add_table(
    ["ID", "Configuration"],
    [
        ["A1", "LoRA + late-block features + focal"],
        ["A2", "LoRA + multi-scale + focal"],
        ["A3", "LoRA + multi-scale + focal + ordinal + referable"],
    ],
)

add_h("2.4 Metrics", 2)
bullets([
    "Accuracy, macro-F1, QWK (primary), referable accuracy/AUROC",
    "Confusion matrices, ROC curves, McNemar tests, bootstrap ΔQWK CIs",
])

add_h("3. Results")
add_h("3.1 Overall performance (APTOS + Messidor-2)", 2)
add_table(
    ["Model", "Set", "Acc", "Macro-F1", "QWK", "Ref Acc", "Ref AUROC"],
    [
        ["RETFound Baseline", "APTOS-test", "0.831", "0.671", "0.888", "0.929", "0.981"],
        ["Enhanced RETFound", "APTOS-test", "0.811", "0.683", "0.893", "0.935", "0.987"],
        ["RETFound Baseline", "Messidor-2", "0.612", "0.330", "0.488", "0.807", "0.725"],
        ["Enhanced RETFound", "Messidor-2", "0.593", "0.349", "0.505", "0.802", "0.778"],
    ],
)
bullets([
    "APTOS: Enhanced improves QWK and referable AUROC slightly; Baseline wins accuracy.",
    "Messidor-2: large domain shift for both. Enhanced stronger on QWK (+0.017) and referable AUROC (+0.053).",
    "External gains are more practically meaningful than the tiny APTOS gaps.",
])

add_h("3.2 APTOS per-class behaviour", 2)
add_mixed(
    "Enhanced RETFound improved **recall on mild (grade 1)** and **severe (grade 3)** relative to Baseline, "
    "with reduced recall on moderate (grade 2). Both models remained excellent on grade 0. This matches an "
    "ordinal/imbalance-aware error shift rather than uniform accuracy gains."
)

add_h("3.3 Ablation study (APTOS test)", 2)
add_table(
    ["Variant", "Acc", "Macro-F1", "QWK", "Ref AUROC"],
    [
        ["RETFound Baseline", "0.831", "0.671", "0.888", "0.981"],
        ["Enhanced RETFound (loaded)", "0.811", "0.683", "0.893", "0.987"],
        ["A1 LoRA + late + focal", "0.738", "0.613", "0.892", "0.985"],
        ["A2 LoRA + multi-scale + focal", "0.765", "0.632", "0.895", "0.987"],
        ["A3 + ordinal + referable", "0.764", "0.629", "0.881", "0.988"],
    ],
)
bullets([
    "A1: LoRA + focal already nearly matches Enhanced QWK — PEFT carries most grading signal.",
    "A2: multi-scale yields the best QWK (0.895) and better accuracy than A1.",
    "A3: highest referable AUROC, but lower QWK than A2 under the 15-epoch schedule.",
])

add_h("3.4 Statistical notes", 2)
add_mixed(
    "In the primary APTOS comparison, bootstrap 95% CI on ΔQWK included 0 and McNemar on referable decisions "
    "was non-significant. Internal edges should be described as modest/not conclusively significant, while "
    "external referable AUROC provides stronger supporting evidence."
)

add_h("4. Discussion")
add_h("4.1 Main findings", 2)
for i, item in enumerate([
    "RETFound Baseline is already strong on APTOS (QWK ≈ 0.89).",
    "Enhanced RETFound is comparable on APTOS and more favourable under Messidor-2 shift for QWK and referable discrimination.",
    "Ablations identify A2 (LoRA + multi-scale + focal) as the best pure grading recipe on APTOS; full Enhanced remains attractive when referable AUROC is prioritised.",
    "Accuracy alone would favour Baseline; ordinal and screening metrics favour Enhanced/A2 — justifying multi-metric reporting.",
], 1):
    p = doc.add_paragraph(f"{i}. {item}")
    for run in p.runs:
        set_run_font(run)

add_h("4.2 Clinical implication", 2)
add_mixed(
    "If the goal is **screening referable disease across sites/cameras**, Enhanced RETFound’s Messidor-2 "
    "referable AUROC gain (0.725 → 0.778) is the most relevant result. If the goal is **maximum in-domain "
    "accuracy on APTOS-like data**, Baseline remains competitive."
)

add_h("4.3 Limitations", 2)
bullets([
    "Single external dataset; absolute Messidor-2 QWK remains modest (~0.50).",
    "Ablations used fewer epochs than the primary Enhanced run.",
    "APTOS ΔQWK was small and not statistically conclusive.",
    "No multi-seed variance reported in this write-up.",
])

add_h("4.4 Future work", 2)
bullets([
    "Longer A3 training / loss-weight search",
    "Additional externals (DDR, EyePACS subsets)",
    "Multi-seed mean ± std",
    "Calibration and decision-curve analysis",
])

add_h("5. Conclusion")
add_mixed(
    "We evaluated RETFound Baseline against Enhanced RETFound for diabetic retinopathy grading, with ablations "
    "and Messidor-2 external testing. Enhanced RETFound matched or slightly exceeded the baseline on APTOS "
    "ordinal/screening metrics and showed clearer external gains in QWK and referable AUROC. Ablations indicate "
    "that LoRA plus multi-scale focal training is the key grading ingredient, while ordinal/referable auxiliaries "
    "mainly boost referable discrimination. Clinically oriented adaptation of RETFound is warranted for "
    "cross-dataset screening performance, even when in-domain accuracy gains are limited."
)

add_h("Figures to include")
bullets([
    "Fig 1: Validation loss/QWK curves (Baseline, Enhanced, A1–A3)",
    "Fig 2–3: APTOS confusion matrices (Baseline vs Enhanced)",
    "Fig 4: APTOS referable ROC (and optional all-models overlay)",
    "Fig 5–6: Messidor-2 confusion matrices",
    "Fig 7: Messidor-2 referable ROC",
])

add_h("Appendix — One-sentence takeaway")
add_mixed(
    "Enhanced RETFound does not dominate APTOS accuracy, but improves ordinal/screening behaviour and "
    "generalises better to Messidor-2 referable detection; multi-scale LoRA fine-tuning is the main useful enhancement."
)

doc.save(out_path)
print("Saved:", out_path)
