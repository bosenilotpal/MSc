"""Rebuild thesis report in IEEE Access-like style/format."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "figures"
DOCX = ROOT / "manuscript" / "Thesis_Report_RETFound_Final.docx"
try:
    test_lock = open(DOCX, "a"); test_lock.close()
except Exception:
    DOCX = ROOT / "manuscript" / "Thesis_Report_RETFound_Final_v2.docx"

BLACK = RGBColor(0, 0, 0)
BODY_PT = 10
ABS_PT = 9
CAPTION_PT = 8


def set_run_font(run, bold=False, italic=False, size=BODY_PT, color=BLACK):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def force_black(run):
    set_run_font(run, bold=run.bold, italic=run.italic,
                 size=run.font.size.pt if run.font.size else BODY_PT)
    rPr = run._element.get_or_add_rPr()
    for child in list(rPr):
        if child.tag == qn("w:color"):
            rPr.remove(child)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rPr.append(color)


def set_cols(section, num=1, space_twips=720):
    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn("w:cols"):
            sectPr.remove(child)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"), str(num))
    if num > 1:
        cols.set(qn("w:space"), str(space_twips))
    sectPr.append(cols)


def add_center(doc, text, bold=False, italic=False, size=10, space_after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic, size=size)
    return p


def add_mixed(doc, text, size=BODY_PT, first_indent=True, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(0.2)
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True, size=size)
        else:
            run = p.add_run(part)
            set_run_font(run, size=size)
    return p


def add_section_h(doc, text):
    """I. INTRODUCTION style — all caps, bold, black, left."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run(text.upper())
    set_run_font(run, bold=True, size=BODY_PT)
    force_black(run)
    return p


def add_sub_h(doc, text):
    """A. Dataset style — bold, black (IEEE often italic for subheads; bold is clearer in Word)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run(text)
    set_run_font(run, bold=True, italic=True, size=BODY_PT)
    force_black(run)
    return p


def add_label_h(doc, text):
    """ABSTRACT / INDEX TERMS — bold centered-left all caps."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    set_run_font(run, bold=True, size=BODY_PT)
    force_black(run)
    return p


def shade(cell, color="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_ieee_table(doc, roman, title, headers, rows):
    # TABLE I
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"TABLE {roman}")
    set_run_font(r, bold=True, size=CAPTION_PT)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run(title)
    set_run_font(r2, italic=True, size=CAPTION_PT)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=8)
        shade(cell, "E7E6E6")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_figure(doc, filename, caption, width=3.2):
    path = FIG / filename
    if not path.exists():
        print("MISSING:", filename)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_after = Pt(8)
    # IEEE: "Fig. 1. Caption text."
    r = cap.add_run(caption)
    set_run_font(r, size=CAPTION_PT)
    print("embedded:", filename)


def add_two_figures(doc, f1, c1, f2, c2, width=1.55):
    table = doc.add_table(rows=2, cols=2)
    for col, (fn, cap) in enumerate([(f1, c1), (f2, c2)]):
        path = FIG / fn
        cell0 = table.rows[0].cells[col]
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path.exists():
            cell0.paragraphs[0].add_run().add_picture(str(path), width=Inches(width))
            print("embedded:", fn)
        else:
            print("MISSING:", fn)
        cell1 = table.rows[1].cells[col]
        cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell1.paragraphs[0].add_run(cap)
        set_run_font(r, size=CAPTION_PT)
    doc.add_paragraph()


def add_ref(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=8)


# ===================== BUILD =====================
doc = Document()

# IEEE Access-like page setup
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    set_cols(section, 1)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(BODY_PT)
style.font.color.rgb = BLACK
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.08
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# --- Front matter (single column) ---
add_center(doc, "Received (date of submission); accepted (date of acceptance); date of publication (date of publication)",
           size=8, space_after=0)
add_center(doc, "Digital Object Identifier 10.1109/ACCESS.2025.XXXXXXX", size=8, space_after=2)
add_center(doc, "VOLUME XX, 2025", size=8, space_after=8)

add_center(
    doc,
    "LoRA-Adapted RETFound with Multi-Scale Fusion and Ordinal Losses\nfor Five-Class Diabetic Retinopathy Grading and External Validation",
    bold=True, size=18, space_after=8,
)
add_center(doc, "NILOTPAL BOSE", bold=True, size=11, space_after=2)
add_center(
    doc,
    "Department of Computer Science and Engineering (update affiliation as needed)",
    italic=True, size=9, space_after=1,
)
add_center(
    doc,
    "Corresponding author: Nilotpal Bose (e-mail: update@email.com)",
    size=8, space_after=8,
)

add_label_h(doc, "Abstract")
add_mixed(
    doc,
    "Adapting retinal foundation models for ordered DR severity remains challenging when flat cross-entropy ignores "
    "class imbalance and clinical ordinal structure. We compare RETFound Baseline (full fine-tuning with weighted "
    "cross-entropy) and Enhanced RETFound (LoRA adapters, multi-scale token fusion, and focal/ordinal/referable losses) "
    "for five-class diabetic retinopathy grading. On APTOS 2019 (test N = 550), Enhanced RETFound slightly improved "
    "quadratic weighted kappa (QWK 0.893 vs 0.888) and referable AUROC (0.987 vs 0.981), while the baseline retained "
    "higher accuracy (0.831 vs 0.811). On external Messidor-2 without fine-tuning, Enhanced RETFound improved QWK "
    "(0.505 vs 0.488) and referable AUROC (0.778 vs 0.725). Ablations show A2 (LoRA + multi-scale + focal) achieves "
    "the best APTOS QWK (0.895), whereas ordinal/referable auxiliaries mainly boost screening AUROC. Overall, "
    "Enhanced RETFound is comparable in-domain and stronger under external shift for ordinal/screening metrics.",
    size=ABS_PT, first_indent=False, space_after=6,
)

add_label_h(doc, "Index Terms")
add_mixed(
    doc,
    "Diabetic retinopathy, RETFound, foundation model, LoRA, multi-scale fusion, ordinal regression, "
    "focal loss, quadratic weighted kappa, Messidor-2, APTOS 2019, transfer learning, medical image classification.",
    size=ABS_PT, first_indent=False, space_after=8,
)

# --- Body: continuous two-column section ---
body = doc.add_section(WD_SECTION.CONTINUOUS)
body.top_margin = Inches(0.75)
body.bottom_margin = Inches(0.75)
body.left_margin = Inches(0.75)
body.right_margin = Inches(0.75)
set_cols(body, 2, space_twips=360)

add_section_h(doc, "I. Introduction")
add_mixed(
    doc,
    "Diabetic retinopathy is a microvascular complication of diabetes and a major cause of preventable vision loss "
    "worldwide. Population screening relies on grading colour fundus photographs according to clinical severity scales "
    "such as the International Clinical Diabetic Retinopathy (ICDR) system, which orders disease from absent (grade 0) "
    "through mild, moderate, and severe non-proliferative disease to proliferative DR (grade 4). Because grading "
    "is labour-intensive and subject to inter-observer variability, deep learning has been widely explored for "
    "automated DR detection and staging.",
)
add_mixed(
    doc,
    "Early systems were typically trained from scratch or from ImageNet-initialised convolutional networks on large "
    "labelled fundus corpora. More recently, foundation models pretrained with self-supervised learning on massive "
    "unlabelled medical image collections have emerged as a stronger starting point for downstream clinical tasks. "
    "**RETFound** is a retinal foundation model based on a Vision Transformer pretrained with masked autoencoding "
    "on approximately 1.6 million fundus and OCT images, and has demonstrated strong transfer to ocular and systemic "
    "prediction tasks with limited labels.",
)
add_mixed(
    doc,
    "Despite this progress, adapting RETFound to DR severity grading still raises open questions. Standard full "
    "fine-tuning with flat cross-entropy treats grades as unordered classes and may under-emphasise rare mild/severe "
    "cases and ordinal structure. Parameter-efficient fine-tuning methods such as LoRA, multi-scale feature "
    "aggregation, and imbalance-/ordinal-aware losses (e.g., focal loss and CORAL-style ordinal regression) "
    "offer a more clinically aligned adaptation recipe. In addition, strong in-domain scores on a single public dataset "
    "(e.g., APTOS 2019) do not guarantee robustness under camera and population shift, motivating external "
    "evaluation on Messidor-2.",
)
add_mixed(
    doc,
    "The principal contributions of this work are: (1) a direct comparison of RETFound Baseline versus Enhanced "
    "RETFound for five-class DR grading; (2) Messidor-2 external validation without fine-tuning; (3) ablations "
    "isolating LoRA/focal, multi-scale fusion, and ordinal/referable auxiliaries; and (4) emphasis on QWK and "
    "referable AUROC alongside accuracy.",
)

add_section_h(doc, "II. Materials and Method")
add_sub_h(doc, "A. Dataset Description")
add_mixed(
    doc,
    "**APTOS 2019 Blindness Detection** provides labelled macular-centred fundus photographs collected in clinical "
    "settings in India. Each image is annotated with an ICDR-aligned severity grade from 0 to 4. The dataset is publicly "
    "known for class imbalance (no-DR and moderate grades dominate; mild and severe grades are scarce), which makes "
    "macro-averaged and ordinal metrics especially informative. We used a stratified 70%/15%/15% train/validation/test "
    "split fixed in advance (test N = 550). Images were resized to 224×224, normalised with ImageNet statistics, and "
    "lightly augmented during training (flips, small rotations, mild colour jitter).",
)
add_mixed(
    doc,
    "**Messidor-2** is used strictly as an **external test set** (no fine-tuning). It contains fundus examinations "
    "acquired under different devices and protocols from APTOS, and therefore probes domain generalisation. We evaluate "
    "five-class grading and a clinically actionable binary endpoint: **referable DR**, defined as grade ≥ 2 (moderate "
    "or worse), consistent with common screening practice.",
)

add_sub_h(doc, "B. Frameworks for Comparison")
add_mixed(
    doc,
    "Two main frameworks share the same RETFound ViT-Large/16 colour-fundus pretrained backbone and the same APTOS "
    "split and preprocessing pipeline.",
)
add_mixed(
    doc,
    "**RETFound Baseline.** The entire encoder is fully fine-tuned with a linear classification head and weighted "
    "cross-entropy. This matches the conventional transfer-learning recipe for foundation-model adaptation and serves "
    "as a strong reference.",
)
add_mixed(
    doc,
    "**Enhanced RETFound.** To reduce catastrophic forgetting and focus learning on DR-relevant adaptation, the backbone "
    "is largely frozen and updated through **LoRA** adapters on attention qkv projections (approximately 1.3% "
    "trainable parameters). Features are taken from multiple transformer depths (blocks 7, 15, and 23) and fused before "
    "classification, capturing both local lesion cues and global severity context. Training optimises a combination of "
    "**focal loss** for hard/rare classes, **CORAL ordinal loss** for ordered severity, and a binary "
    "**referable-DR** auxiliary head. Checkpoints for both frameworks were selected by best validation QWK.",
)
add_mixed(
    doc,
    "**Ablations (A1–A3).** To isolate contributions, we trained three Enhanced variants for 15 epochs on the same split: "
    "**A1** LoRA + late features + focal; **A2** A1 + multi-scale fusion; **A3** A2 + ordinal + referable auxiliaries.",
)
add_mixed(
    doc,
    "**Evaluation.** We report accuracy, macro-F1, QWK, referable accuracy/AUROC, confusion matrices, and ROC "
    "curves. Primary interpretation emphasises QWK for grading agreement and referable AUROC for screening utility.",
)

add_section_h(doc, "III. Results")
add_sub_h(doc, "A. Main Comparison on APTOS")
add_ieee_table(
    doc, "I",
    "Main Results on APTOS and Messidor-2 (checkpoints selected by validation QWK).",
    ["Model", "Set", "Acc", "Macro-F1", "QWK", "Ref Acc", "Ref AUROC"],
    [
        ["RETFound Baseline", "APTOS", "0.831", "0.671", "0.888", "0.929", "0.981"],
        ["Enhanced RETFound", "APTOS", "0.811", "0.683", "0.893", "0.935", "0.987"],
        ["RETFound Baseline", "Messidor-2", "0.612", "0.330", "0.488", "0.807", "0.725"],
        ["Enhanced RETFound", "Messidor-2", "0.593", "0.349", "0.505", "0.802", "0.778"],
    ],
)
add_mixed(
    doc,
    "On the APTOS held-out test set, Enhanced RETFound yields a small gain in QWK (+0.005) and referable AUROC "
    "(+0.006), while Baseline remains higher in overall accuracy (+0.020). This pattern indicates that the enhancement "
    "does not uniformly dominate all metrics; instead, it shifts performance toward ordinal agreement and "
    "screening-oriented discrimination. Because APTOS gains are modest, external validation and ablations are "
    "essential to judge whether the Enhanced design is practically useful.",
)
add_figure(doc, "fig1_aptos_comparison.png",
           "Fig. 1. APTOS test metrics for RETFound Baseline versus Enhanced RETFound.")
add_figure(doc, "fig3_internal_vs_external.png",
           "Fig. 2. Internal (APTOS) versus external (Messidor-2) QWK and referable AUROC.")

add_sub_h(doc, "B. Validation QWK Curves")
add_mixed(
    doc,
    "Fig. 3 shows validation QWK versus training epoch for the two main frameworks. Both models improve rapidly "
    "in early epochs and then plateau; checkpoints used for testing were selected by maximum validation QWK rather "
    "than the final epoch, which reduces the risk of reporting an overfit late checkpoint.",
)
add_two_figures(doc,
    "main_baseline_qwk.png", "Fig. 3. (a) Baseline — validation QWK.",
    "main_enhanced_qwk.png", "(b) Enhanced — validation QWK.")
add_mixed(
    doc,
    "Fig. 4 shows validation QWK trajectories for the ablation variants (15 epochs). These curves help confirm "
    "that A2’s superior test QWK is consistent with a stable validation trend, rather than a single noisy checkpoint.",
)
add_figure(doc, "ablation_A1_lora_late_focal_qwk.png",
           "Fig. 4. (a) A1 — validation QWK.", width=3.0)
add_figure(doc, "ablation_A2_lora_ms_focal_qwk.png",
           "(b) A2 — validation QWK.", width=3.0)
add_figure(doc, "ablation_A3_lora_ms_ord_ref_qwk.png",
           "(c) A3 — validation QWK.", width=3.0)

add_sub_h(doc, "C. External Validation on Messidor-2")
add_mixed(
    doc,
    "External validation tests whether models trained on APTOS remain useful when imaging devices, populations, "
    "and acquisition protocols change. Messidor-2 was therefore evaluated **without any additional fine-tuning**, "
    "using the same preprocessing and the APTOS-selected checkpoints.",
)
add_mixed(
    doc,
    "**Overall external behaviour.** Absolute performance falls for both models relative to APTOS (Baseline QWK "
    "0.888 → 0.488; Enhanced QWK 0.893 → 0.505). Such degradation is expected under domain shift and should not be "
    "interpreted as training failure. The scientifically relevant question is the **relative** gap between frameworks "
    "on the external set.",
)
add_mixed(
    doc,
    "**Where Enhanced RETFound helps externally.** Enhanced RETFound improves Messidor-2 QWK from 0.488 to 0.505 "
    "(+0.017) and referable AUROC from 0.725 to 0.778 (+0.053). The referable-AUROC gain is substantially larger than "
    "the corresponding APTOS gap and is the strongest evidence that clinically oriented adaptation improves "
    "cross-dataset screening behaviour. Macro-F1 also favours Enhanced (0.349 vs 0.330). Baseline retains a small edge "
    "in raw accuracy (0.612 vs 0.593) and referable accuracy (0.807 vs 0.802), again showing that accuracy alone would "
    "understate Enhanced’s external benefit.",
)
add_mixed(
    doc,
    "**Why external metrics matter clinically.** In screening workflows, the priority is often to rank or detect eyes "
    "that need referral (moderate or worse DR), not merely to maximise exact five-class accuracy on the training "
    "distribution. The Messidor-2 referable ROC comparison supports Enhanced RETFound in this role: its ROC lies above "
    "Baseline across operating points, consistent with the AUROC increase.",
)
add_mixed(
    doc,
    "**Error-structure observation.** External confusion matrices show broader grade confusion than on APTOS for both "
    "models, reflecting harder domain transfer. Nevertheless, Enhanced’s better QWK implies that remaining errors are, "
    "on average, more ordinal-consistent (closer grades) than Baseline’s, which is preferable when severity is an "
    "ordered clinical construct.",
)
add_two_figures(doc,
    "cm_baseline_external.png", "Fig. 5. (a) Baseline — Messidor-2 CM.",
    "cm_enhanced_external.png", "(b) Enhanced — Messidor-2 CM.")
add_two_figures(doc,
    "roc_referable_aptos.png", "Fig. 6. (a) Referable ROC — APTOS.",
    "roc_referable_external.png", "(b) Referable ROC — Messidor-2.")

add_sub_h(doc, "D. Confusion Matrices on APTOS")
add_mixed(
    doc,
    "APTOS confusion matrices show both models are excellent on grade 0 (no DR). Enhanced RETFound improves recall "
    "for mild (grade 1) and severe (grade 3) disease relative to Baseline, at the cost of more confusion involving "
    "moderate disease (grade 2). This is consistent with adjacent-grade ordinal errors, which QWK penalises less "
    "harshly than flat accuracy. Referable ROC curves remain strong for both models on APTOS (AUROC > 0.98), so "
    "in-domain screening discrimination is already near ceiling; the more discriminative ROC comparison appears on "
    "Messidor-2 (Section III-C).",
)
add_two_figures(doc,
    "cm_baseline_aptos.png", "Fig. 7. (a) Baseline — APTOS CM.",
    "cm_enhanced_aptos.png", "(b) Enhanced — APTOS CM.")

add_sub_h(doc, "E. Ablation Study")
add_mixed(
    doc,
    "Ablations answer which design choices inside Enhanced RETFound actually matter. All ablation variants were "
    "trained on the **same APTOS split** for 15 epochs, with checkpoints selected by validation QWK, and then "
    "evaluated on the APTOS test set.",
)
add_ieee_table(
    doc, "II",
    "Ablation Results on APTOS (15 epochs; best validation-QWK checkpoint).",
    ["Variant", "What is enabled", "Acc", "QWK", "Ref AUROC"],
    [
        ["RETFound Baseline", "Full FT + CE", "0.831", "0.888", "0.981"],
        ["Enhanced RETFound", "Full Enhanced (primary ckpt)", "0.811", "0.893", "0.987"],
        ["A1", "LoRA + late features + focal", "0.738", "0.892", "0.985"],
        ["A2", "A1 + multi-scale fusion", "0.765", "0.895", "0.987"],
        ["A3", "A2 + ordinal + referable losses", "0.764", "0.881", "0.988"],
    ],
)
add_mixed(
    doc,
    "**A1 — LoRA + focal only.** A1 already reaches QWK 0.892, nearly matching the full Enhanced checkpoint (0.893). "
    "This shows that parameter-efficient adaptation of RETFound, combined with class-imbalance-aware focal loss, "
    "accounts for most of the ordinal grading signal. Accuracy is lower than Baseline (0.738 vs 0.831), indicating "
    "A1 redistributes errors rather than simply fitting majority classes. In other words, LoRA is not a weak "
    "substitute here; it is the core mechanism of the Enhanced framework.",
)
add_mixed(
    doc,
    "**A2 — adding multi-scale fusion.** Adding multi-scale token fusion improves QWK to 0.895 (best among all "
    "variants) and raises accuracy from 0.738 to 0.765 versus A1. Referable AUROC also matches the Enhanced model "
    "(0.987). This supports the hypothesis that DR grading benefits from combining mid-level local cues "
    "(microaneurysms/exudates) with deeper global context (overall severity). Among the Enhanced family, A2 is the "
    "strongest pure grading recipe.",
)
add_mixed(
    doc,
    "**A3 — adding ordinal and referable auxiliaries.** A3 attains the highest referable AUROC (0.988) but reduces "
    "QWK to 0.881, below both A2 and Baseline. Two interpretations are compatible with the data: (i) the extra loss "
    "terms emphasise screening separation (referable vs non-referable) more than fine-grained ordinal agreement; "
    "(ii) under a short 15-epoch ablation schedule, the multi-objective loss may be under-optimised relative to A2. "
    "Practically, A3 is attractive if referable detection is the deployment goal, whereas A2 is preferable if "
    "five-class QWK is the primary endpoint. The primary Enhanced model, trained longer with the full objective, "
    "sits between these behaviours (QWK 0.893, AUROC 0.987).",
)
add_mixed(
    doc,
    "**Summary of ablation implications.** (1) LoRA + focal (A1) is necessary and nearly sufficient for "
    "Enhanced-level QWK. (2) Multi-scale fusion (A2) provides the clearest grading gain. (3) Ordinal/referable "
    "auxiliaries (A3) help screening AUROC but can hurt QWK if not carefully scheduled/weighted. Therefore, the "
    "Enhanced design is justified as a configurable clinical adaptation stack rather than a single monolithic architecture.",
)
add_figure(doc, "fig4_ablation_qwk_auroc.png",
           "Fig. 8. Ablation summary on APTOS: QWK (left) and referable AUROC (right). "
           "A2 achieves the highest QWK; A3 achieves the highest referable AUROC.")

add_section_h(doc, "IV. Conclusion")
add_mixed(
    doc,
    "This work compared standard RETFound fine-tuning with a clinically motivated Enhanced RETFound recipe for DR "
    "severity grading. On APTOS, the enhanced model is comparable to the baseline, with small gains in QWK and "
    "referable AUROC but lower accuracy. External Messidor-2 testing provides stronger support: Enhanced RETFound "
    "improves QWK and, especially, referable AUROC under domain shift, which matters for screening-style deployment. "
    "Ablations show that LoRA + focal adaptation already recovers most Enhanced-level QWK, multi-scale fusion (A2) "
    "gives the best grading agreement, and ordinal/referable auxiliaries (A3) mainly boost screening AUROC. "
    "Foundation-model adaptation for DR should therefore be judged with ordinal and clinical endpoints—not accuracy "
    "alone—and external validation is essential before claiming practical benefit.",
)

# References often better in single column for readability — new continuous section
refs_sec = doc.add_section(WD_SECTION.CONTINUOUS)
refs_sec.top_margin = Inches(0.75)
refs_sec.bottom_margin = Inches(0.75)
refs_sec.left_margin = Inches(0.75)
refs_sec.right_margin = Inches(0.75)
set_cols(refs_sec, 1)

add_section_h(doc, "References")
refs = [
    "[1] Y. Zhou et al., “A foundation model for generalizable disease detection from retinal images,” Nature, vol. 622, pp. 156–163, 2023.",
    "[2] A. Dosovitskiy et al., “An image is worth 16×16 words: Transformers for image recognition at scale,” in Proc. ICLR, 2021.",
    "[3] K. He et al., “Masked autoencoders are scalable vision learners,” in Proc. CVPR, 2022, pp. 16000–16009.",
    "[4] E. J. Hu et al., “LoRA: Low-rank adaptation of large language models,” in Proc. ICLR, 2022.",
    "[5] T.-Y. Lin, P. Goyal, R. Girshick, K. He, and P. Dollár, “Focal loss for dense object detection,” in Proc. ICCV, 2017, pp. 2980–2988.",
    "[6] W. Cao, V. Mirjalili, and S. Raschka, “Rank consistent ordinal regression for neural networks with application to age estimation,” Pattern Recognit. Lett., vol. 140, pp. 325–331, 2020.",
    "[7] C. P. Wilkinson et al., “Proposed international clinical diabetic retinopathy and diabetic macular edema disease severity scales,” Ophthalmology, vol. 110, no. 9, pp. 1677–1682, 2003.",
    "[8] V. Gulshan et al., “Development and validation of a deep learning algorithm for detection of diabetic retinopathy in retinal fundus photographs,” JAMA, vol. 316, no. 22, pp. 2402–2410, 2016.",
    "[9] E. Decencière et al., “Feedback on a publicly distributed image database: the Messidor database,” Image Anal. Stereol., vol. 33, no. 3, pp. 231–234, 2014.",
    "[10] Karthik, Maggie, and S. Dane, “APTOS 2019 Blindness Detection,” Kaggle, 2019. [Online]. Available: https://www.kaggle.com/competitions/aptos2019-blindness-detection",
    "[11] J. Krause et al., “Grader variability and the importance of reference standards for evaluating machine learning models for diabetic retinopathy,” Ophthalmology, vol. 125, no. 8, pp. 1264–1272, 2018.",
    "[12] D. S. W. Ting et al., “Artificial intelligence and deep learning in ophthalmology,” Br. J. Ophthalmol., vol. 103, no. 2, pp. 167–175, 2019.",
    "[13] T. Li et al., “Diagnostic assessment of deep learning algorithms for diabetic retinopathy screening,” Inf. Sci., vol. 501, pp. 511–522, 2019.",
    "[14] J. Cohen, “Weighted kappa: Nominal scale agreement with provision for scaled disagreement or partial credit,” Psychol. Bull., vol. 70, no. 4, pp. 213–220, 1968.",
    "[15] R. Bommasani et al., “On the opportunities and risks of foundation models,” arXiv:2108.07258, 2021.",
]
for r in refs:
    add_ref(doc, r)

doc.save(DOCX)
print("Saved IEEE-style Word:", DOCX)
