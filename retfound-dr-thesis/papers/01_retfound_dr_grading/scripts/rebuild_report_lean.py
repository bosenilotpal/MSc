"""Rebuild a lean Word thesis report with only the most important figures."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "figures"
DOCX = ROOT / "manuscript" / "Thesis_Report_RETFound_Baseline_vs_Enhanced.docx"


def set_run_font(run, bold=False, italic=False, size=12):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_mixed(doc, text, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2]); set_run_font(run, bold=True)
        else:
            run = p.add_run(part); set_run_font(run)


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, size=16 if level == 1 else 13)


def shade(cell, color="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=9)
        shade(cell)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=9)
    doc.add_paragraph()


def add_figure(doc, filename, caption, width=5.6):
    path = FIG / filename
    if not path.exists():
        print("MISSING:", filename)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, italic=True, size=10)
    cap.paragraph_format.space_after = Pt(12)
    print("embedded:", filename)


def add_two_figures(doc, f1, c1, f2, c2, width=3.05):
    table = doc.add_table(rows=2, cols=2)
    for col, (fn, cap) in enumerate([(f1, c1), (f2, c2)]):
        path = FIG / fn
        cell0 = table.rows[0].cells[col]
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path.exists():
            cell0.paragraphs[0].add_run().add_picture(str(path), width=Inches(width))
            print("embedded:", fn)
        else:
            print("MISSING:", fn)
        cell1 = table.rows[1].cells[col]
        cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell1.paragraphs[0].add_run(cap)
        set_run_font(r, italic=True, size=9)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run)


doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.15

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("Enhancing RETFound for Diabetic Retinopathy Grading:\nBaseline vs Clinically Adapted Fine-Tuning")
set_run_font(run, bold=True, size=16)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = s.add_run("Thesis Experimental Report")
set_run_font(run, italic=True, size=12)
doc.add_paragraph()

add_h(doc, "Abstract")
add_mixed(
    doc,
    "We compare **RETFound Baseline** (full fine-tuning + cross-entropy) with **Enhanced RETFound** "
    "(LoRA, multi-scale fusion, and focal/ordinal/referable losses) for five-class DR grading. On APTOS 2019 "
    "(test N = 550), Enhanced RETFound slightly improved QWK (0.893 vs 0.888) and referable AUROC "
    "(0.987 vs 0.981), while the baseline kept higher accuracy (0.831 vs 0.811). On external Messidor-2, "
    "Enhanced RETFound improved QWK (0.505 vs 0.488) and referable AUROC (0.778 vs 0.725). Ablation A2 "
    "(LoRA + multi-scale + focal) achieved the best APTOS QWK (0.895). Enhanced RETFound is comparable "
    "internally and stronger on external referable discrimination."
)

add_h(doc, "1. Introduction")
add_mixed(
    doc,
    "Diabetic retinopathy grading is ordinal and class-imbalanced. This study asks whether a clinically "
    "motivated adaptation of RETFound improves severity grading over standard fine-tuning, including under "
    "external domain shift to Messidor-2."
)

add_h(doc, "2. Methods")
add_table(doc, ["Dataset", "Role"], [
    ["APTOS 2019", "Train/val/test; test N = 550"],
    ["Messidor-2", "External test only"],
])
add_table(doc, ["Model", "Description"], [
    ["RETFound Baseline", "Full fine-tuning + weighted CE"],
    ["Enhanced RETFound", "LoRA + multi-scale + focal/ordinal/referable"],
    ["A1 / A2 / A3", "Ablations of Enhanced components (15 epochs)"],
])
bullets(doc, [
    "Primary metric: Quadratic Weighted Kappa (QWK).",
    "Secondary: accuracy, macro-F1, referable DR accuracy/AUROC.",
])

add_h(doc, "3. Results")

add_h(doc, "3.1 Main comparison", 2)
add_table(doc, ["Model", "Set", "Acc", "Macro-F1", "QWK", "Ref Acc", "Ref AUROC"], [
    ["RETFound Baseline", "APTOS", "0.831", "0.671", "0.888", "0.929", "0.981"],
    ["Enhanced RETFound", "APTOS", "0.811", "0.683", "0.893", "0.935", "0.987"],
    ["RETFound Baseline", "Messidor-2", "0.612", "0.330", "0.488", "0.807", "0.725"],
    ["Enhanced RETFound", "Messidor-2", "0.593", "0.349", "0.505", "0.802", "0.778"],
])
add_figure(doc, "fig1_aptos_comparison.png",
           "Figure 1. APTOS test metrics: RETFound Baseline vs Enhanced RETFound.")
add_figure(doc, "fig3_internal_vs_external.png",
           "Figure 2. Internal vs external generalization (QWK and referable AUROC). "
           "Enhanced RETFound shows a clearer Messidor-2 gain in referable AUROC (0.725 → 0.778).")

add_h(doc, "3.2 Confusion matrices", 2)
add_two_figures(
    doc,
    "cm_baseline_aptos.png", "Figure 3a. Baseline — APTOS",
    "cm_enhanced_aptos.png", "Figure 3b. Enhanced — APTOS",
)
add_two_figures(
    doc,
    "cm_baseline_external.png", "Figure 4a. Baseline — Messidor-2",
    "cm_enhanced_external.png", "Figure 4b. Enhanced — Messidor-2",
)

add_h(doc, "3.3 Referable-DR ROC", 2)
add_two_figures(
    doc,
    "roc_referable_aptos.png", "Figure 5a. Referable ROC — APTOS",
    "roc_referable_external.png", "Figure 5b. Referable ROC — Messidor-2",
)

add_h(doc, "3.4 Ablations", 2)
add_table(doc, ["Variant", "Acc", "QWK", "Ref AUROC"], [
    ["RETFound Baseline", "0.831", "0.888", "0.981"],
    ["Enhanced RETFound", "0.811", "0.893", "0.987"],
    ["A1 LoRA + focal", "0.738", "0.892", "0.985"],
    ["A2 + multi-scale", "0.765", "0.895", "0.987"],
    ["A3 + ordinal/referable", "0.764", "0.881", "0.988"],
])
add_figure(doc, "fig4_ablation_qwk_auroc.png",
           "Figure 6. Ablation summary on APTOS: QWK (left) and referable AUROC (right). "
           "A2 best QWK; A3 best referable AUROC.")

add_h(doc, "4. Discussion")
for i, item in enumerate([
    "Baseline is already strong on APTOS; Enhanced is comparable in-domain.",
    "Under Messidor-2 shift, Enhanced improves QWK and especially referable AUROC.",
    "A2 (LoRA + multi-scale + focal) is the strongest grading configuration among ablations.",
    "Accuracy favours Baseline; ordinal/screening metrics favour Enhanced/A2 — report both.",
], 1):
    p = doc.add_paragraph(f"{i}. {item}")
    for run in p.runs:
        set_run_font(run)

add_h(doc, "5. Conclusion")
add_mixed(
    doc,
    "Enhanced RETFound does not dominate APTOS accuracy, but improves ordinal/screening behaviour and "
    "generalises better to Messidor-2 referable detection. Multi-scale LoRA fine-tuning is the main useful enhancement."
)

doc.save(DOCX)
print("\nSaved lean report:", DOCX)
