"""Rebuild Word thesis report with all available figures from figures/."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

ROOT = Path(__file__).resolve().parents[1]
FIG = ROOT / "figures"
DOCX = ROOT / "manuscript" / "Thesis_Report_RETFound_Baseline_vs_Enhanced.docx"


def set_run_font(run, bold=False, italic=False, size=12):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_mixed(doc, text, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2]); set_run_font(run, bold=True)
        else:
            run = p.add_run(part); set_run_font(run)


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, bold=True, size=16 if level == 1 else 13)


def shade(cell, color="D9E2F3"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, bold=True, size=9)
        shade(cell)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=9)
    doc.add_paragraph()


def add_figure(doc, filename, caption, width=5.8):
    path = FIG / filename
    if not path.exists():
        p = doc.add_paragraph()
        r = p.add_run(f"[Missing figure: {filename}]")
        set_run_font(r, italic=True, size=10)
        print("MISSING:", filename)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, italic=True, size=10)
    cap.paragraph_format.space_after = Pt(12)
    print("embedded:", filename)


def add_two_figures(doc, f1, c1, f2, c2, width=3.1):
    """Place two images side by side in a 1x2 table."""
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True
    for col, (fn, cap) in enumerate([(f1, c1), (f2, c2)]):
        path = FIG / fn
        cell0 = table.rows[0].cells[col]
        cell0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path.exists():
            run = cell0.paragraphs[0].add_run()
            run.add_picture(str(path), width=Inches(width))
            print("embedded:", fn)
        else:
            cell0.paragraphs[0].add_run(f"[Missing: {fn}]")
            print("MISSING:", fn)
        cell1 = table.rows[1].cells[col]
        cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell1.paragraphs[0].add_run(cap)
        set_run_font(r, italic=True, size=9)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run)


doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.15

# Title
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("Enhancing RETFound for Diabetic Retinopathy Grading:\nBaseline vs Clinically Adapted Fine-Tuning")
set_run_font(run, bold=True, size=16)
s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = s.add_run("Thesis Experimental Report (APTOS + Messidor-2 + Ablations)")
set_run_font(run, italic=True, size=12)
doc.add_paragraph()

add_h(doc, "Abstract")
add_mixed(
    doc,
    "We compare **RETFound Baseline** (full fine-tuning + cross-entropy) with **Enhanced RETFound** "
    "(LoRA, multi-scale fusion, and focal/ordinal/referable losses) for five-class DR grading. On APTOS 2019 "
    "(test N = 550), Enhanced RETFound slightly improved QWK (0.893 vs 0.888) and referable AUROC "
    "(0.987 vs 0.981), while the baseline kept higher accuracy (0.831 vs 0.811). On external Messidor-2, "
    "Enhanced RETFound improved QWK (0.505 vs 0.488) and referable AUROC (0.778 vs 0.725) despite lower "
    "accuracy. Ablations showed LoRA + multi-scale + focal (A2) achieved the best APTOS QWK (0.895); adding "
    "ordinal/referable losses raised referable AUROC but reduced QWK under a short training schedule. "
    "Enhanced RETFound is comparable internally and stronger on external referable discrimination."
)

add_h(doc, "1. Introduction")
add_mixed(
    doc,
    "Diabetic retinopathy (DR) grading is **ordinal** (ICDR 0–4) and **class-imbalanced**. RETFound provides "
    "strong retinal representations, but standard full fine-tuning with flat cross-entropy may under-emphasise "
    "rare grades and ordinal structure. This study evaluates whether LoRA + multi-scale + ordinal/focal "
    "adaptation improves grading over standard RETFound fine-tuning, including under external domain shift."
)

add_h(doc, "2. Methods")
add_h(doc, "2.1 Datasets", 2)
add_table(doc, ["Dataset", "Role"], [
    ["APTOS 2019", "Train/val/test (70/15/15; fixed splits); test N = 550"],
    ["Messidor-2", "External test only (no fine-tuning)"],
])
bullets(doc, [
    "Tasks: five-class ICDR grading; secondary endpoint referable DR (grade ≥ 2).",
    "Input: 224×224 fundus images; ImageNet normalisation; light train-time augmentation.",
])

add_h(doc, "2.2 Models", 2)
add_table(doc, ["Model", "Description"], [
    ["RETFound Baseline", "RETFound ViT-L/16 (CFP); full fine-tuning; weighted cross-entropy"],
    ["Enhanced RETFound", "LoRA on qkv; multi-scale blocks 7/15/23; focal + CORAL + referable losses"],
])

add_h(doc, "2.3 Ablations (APTOS, 15 epochs)", 2)
add_table(doc, ["ID", "Configuration"], [
    ["A1", "LoRA + late-block features + focal"],
    ["A2", "LoRA + multi-scale + focal"],
    ["A3", "LoRA + multi-scale + focal + ordinal + referable"],
])

add_h(doc, "2.4 Metrics", 2)
bullets(doc, [
    "Accuracy, macro-F1, QWK (primary), referable accuracy/AUROC",
    "Confusion matrices, ROC curves, McNemar / bootstrap analyses",
])

# ---------------- RESULTS ----------------
add_h(doc, "3. Results")

add_h(doc, "3.1 Overall performance (APTOS + Messidor-2)", 2)
add_table(doc, ["Model", "Set", "Acc", "Macro-F1", "QWK", "Ref Acc", "Ref AUROC"], [
    ["RETFound Baseline", "APTOS-test", "0.831", "0.671", "0.888", "0.929", "0.981"],
    ["Enhanced RETFound", "APTOS-test", "0.811", "0.683", "0.893", "0.935", "0.987"],
    ["RETFound Baseline", "Messidor-2", "0.612", "0.330", "0.488", "0.807", "0.725"],
    ["Enhanced RETFound", "Messidor-2", "0.593", "0.349", "0.505", "0.802", "0.778"],
])

add_figure(doc, "fig1_aptos_comparison.png",
           "Figure 1. APTOS held-out test metrics for RETFound Baseline vs Enhanced RETFound.")
add_figure(doc, "fig2_messidor_comparison.png",
           "Figure 2. Messidor-2 external test metrics for RETFound Baseline vs Enhanced RETFound.")
add_figure(doc, "fig3_internal_vs_external.png",
           "Figure 3. Internal (APTOS) versus external (Messidor-2) QWK and referable AUROC. "
           "Enhanced RETFound shows a clearer external gain in referable discrimination.")

bullets(doc, [
    "APTOS: Enhanced improves QWK and referable AUROC slightly; Baseline wins accuracy.",
    "Messidor-2: Enhanced stronger on QWK (+0.017) and referable AUROC (+0.053).",
])

add_h(doc, "3.2 Training loss curves", 2)
add_mixed(doc, "Training dynamics for the two main models (from saved training histories):")
add_two_figures(
    doc,
    "main_baseline_loss.png", "Figure 4a. RETFound Baseline — training loss.",
    "main_enhanced_loss.png", "Figure 4b. Enhanced RETFound — training loss.",
)
add_mixed(doc, "Ablation training loss curves (15 epochs):")
add_figure(doc, "ablation_A1_lora_late_focal_loss.png",
           "Figure 5. A1 (LoRA + late features + focal) — train/val loss.")
add_figure(doc, "ablation_A2_lora_ms_focal_loss.png",
           "Figure 6. A2 (LoRA + multi-scale + focal) — train/val loss.")
add_figure(doc, "ablation_A3_lora_ms_ord_ref_loss.png",
           "Figure 7. A3 (LoRA + multi-scale + ordinal/referable) — train/val loss.")

add_h(doc, "3.3 Confusion matrices", 2)
add_mixed(doc, "APTOS test set:")
add_two_figures(
    doc,
    "cm_baseline_aptos.png", "Figure 8a. RETFound Baseline — APTOS confusion matrix.",
    "cm_enhanced_aptos.png", "Figure 8b. Enhanced RETFound — APTOS confusion matrix.",
)
add_mixed(doc, "Messidor-2 external test:")
add_two_figures(
    doc,
    "cm_baseline_external.png", "Figure 9a. RETFound Baseline — Messidor-2 confusion matrix.",
    "cm_enhanced_external.png", "Figure 9b. Enhanced RETFound — Messidor-2 confusion matrix.",
)

add_h(doc, "3.4 Referable-DR ROC curves", 2)
add_two_figures(
    doc,
    "roc_referable_aptos.png", "Figure 10a. Referable DR ROC — APTOS test.",
    "roc_referable_external.png", "Figure 10b. Referable DR ROC — Messidor-2.",
    width=3.05,
)
add_mixed(
    doc,
    "On Messidor-2, the referable AUROC gap favouring Enhanced RETFound (0.778 vs 0.725) is visually "
    "consistent with the ROC curves above."
)

add_h(doc, "3.5 Ablation study", 2)
add_table(doc, ["Variant", "Acc", "Macro-F1", "QWK", "Ref AUROC"], [
    ["RETFound Baseline", "0.831", "0.671", "0.888", "0.981"],
    ["Enhanced RETFound (loaded)", "0.811", "0.683", "0.893", "0.987"],
    ["A1 LoRA + late + focal", "0.738", "0.613", "0.892", "0.985"],
    ["A2 LoRA + multi-scale + focal", "0.765", "0.632", "0.895", "0.987"],
    ["A3 + ordinal + referable", "0.764", "0.629", "0.881", "0.988"],
])
add_figure(doc, "fig4_ablation_qwk_auroc.png",
           "Figure 11. Ablation QWK (left) and referable AUROC (right) on APTOS. "
           "A2 achieves the highest QWK; A3 achieves the highest referable AUROC.")
add_figure(doc, "fig5_ablation_acc_vs_qwk.png",
           "Figure 12. Ablation trade-off between accuracy and QWK on APTOS.")

bullets(doc, [
    "A1: LoRA + focal already nearly matches Enhanced QWK.",
    "A2: multi-scale yields the best QWK (0.895).",
    "A3: highest referable AUROC, but lower QWK than A2 under the short ablation schedule.",
])

add_h(doc, "4. Discussion")
for i, item in enumerate([
    "RETFound Baseline is already strong on APTOS (QWK ≈ 0.89).",
    "Enhanced RETFound is comparable on APTOS and more favourable under Messidor-2 shift for QWK and referable AUROC.",
    "Confusion matrices and ROCs support a clinically oriented error shift and stronger external referable discrimination for Enhanced RETFound.",
    "Ablations identify A2 (LoRA + multi-scale + focal) as the best pure grading recipe; full Enhanced remains attractive when referable AUROC is prioritised.",
], 1):
    p = doc.add_paragraph(f"{i}. {item}")
    for run in p.runs:
        set_run_font(run)

add_h(doc, "5. Conclusion")
add_mixed(
    doc,
    "Enhanced RETFound matched or slightly exceeded Baseline on APTOS ordinal/screening metrics and showed "
    "clearer Messidor-2 gains in QWK and referable AUROC. Ablations indicate that LoRA plus multi-scale focal "
    "training is the key grading ingredient, while ordinal/referable auxiliaries mainly boost referable discrimination."
)

doc.save(DOCX)
print("\nSaved:", DOCX)
print("Figures used from:", FIG)
